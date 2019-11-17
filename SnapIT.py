from PIL import ImageGrab
import xlsxwriter
import openpyxl
from openpyxl import load_workbook
import time
import os
from os import path
from pathvalidate import is_valid_filename
import tkinter
from tkinter import ttk,filedialog,StringVar,END,Toplevel
from tkinter import Tk as ThemedTk
from docx import Document
from docx.shared import Inches
from ttkthemes import ThemedTk, THEMES
root = ThemedTk()
root.set_theme('scidblue')

import tkinter as tk 
class CollapsiblePane(ttk.Frame): 
    def __init__(self, parent, expanded_text ="Collapse <<", 
                               collapsed_text ="Expand >>"): 
  
        ttk.Frame.__init__(self, parent) 
        self.parent = parent 
        self._expanded_text = expanded_text 
        self._collapsed_text = collapsed_text 
  
        # Here weight implies that it can grow it's 
        # size if extra space is available 
        # default weight is 0 
        self.columnconfigure(1, weight = 1) 
  
        # Tkinter variable storing integer value 
        self._variable = tk.IntVar() 
  
        # Checkbutton is created but will behave as Button 
        # cause in style, Button is passed 
        # main reason to do this is Button do not support 
        # variable option but checkbutton do 
        self._button = ttk.Checkbutton(self,variable = self._variable, 
                            command = self._activate, style ="TButton") 
        self._button.grid(row = 0, columnspan=2,sticky='nwes') 
  
        # This wil create a seperator 
        # A separator is a line, we can also set thickness 
        # self._separator = ttk.Separator(self, orient ="horizontal") 
        # self._separator.grid(row = 0, column = 1, sticky ="nwes") 
  
        self.frame = ttk.Frame(self) 
  
        # This will call activate functon of class 
        self._activate() 
  
    def _activate(self): 
        if  self._variable.get(): 
  
            # As soon as button is pressed it removes this widget 
            # but is not destroyed means can be displayed again 
            self.frame.grid_forget() 
  
            # This will change the text of the checkbutton 
            self._button.configure(text = self._collapsed_text) 
  
        elif not self._variable.get(): 
            # increasing the frame area so new widgets 
            # could reside in this container 
            self.frame.grid(row = 1,columnspan = 2,sticky='news') 
            self._button.configure(text = self._expanded_text) 
  
    def toggle(self): 
        """Switches the label frame to the opposite state."""
        self._variable.set(not self._variable.get()) 
        self._activate()

def disable():
     selectPath.configure(state='disabled')
     Create.configure(state='disabled')
     ScreenShot.configure(state='disabled')
     ClipBoard.configure(state='disabled')
     Comment.configure(state='disabled')
     e.configure(state='disabled')
     combo.configure(state='disabled')

def enable():
     selectPath.configure(state='normal')
     Create.configure(state='normal')
     ScreenShot.configure(state='normal')
     ClipBoard.configure(state='normal')
     Comment.configure(state='normal')
     e.configure(state='normal')
     combo.configure(state='normal')

def Errormessage(ErrorMsg): 
     disable()
     child = Toplevel(root) 
     child.title('Error')
     child.resizable(False,False)
     windowWidth = root.winfo_reqwidth()
     windowHeight = root.winfo_reqheight()
     positionRight = int(root.winfo_screenwidth()/2 - windowWidth/2)
     positionDown = int(root.winfo_screenheight()/2 - windowHeight/2)
     child.geometry("+{}+{}".format(positionRight, positionDown))
     ttk.Label(child,text = ErrorMsg,anchor='center').grid(row=1, sticky='nwes')
     def close_win():
          enable()
          child.destroy()
     ttk.Button(child,text='OK',width=25,command=close_win).grid(row=2,sticky='nwes')
     child.protocol("WM_DELETE_WINDOW", close_win)

def createExcel():
    filename=e.get()
    filepath=str(z.get()).replace("/","\\\\")
    if(path.exists(filepath+'\\'+filename+'.xlsx')):
         Errormessage('File already exists in Path')
    else:
         workbook=xlsxwriter.Workbook(filepath+'\\'+filename+'.xlsx')
         workbook.add_worksheet()
         workbook.close()
         

def createWord():
    filename=e.get()
    filepath=str(z.get()).replace("/","\\\\")
    if(path.exists(filepath+'\\'+filename+'.docx')):
         Errormessage('File already exists in Path')
    else:
         document = Document()
         document.save(filepath+'\\'+filename+'.docx')

def CreateFile():
    filename=e.get()
    filepath=str(z.get()).replace("/","\\\\")
    if(filepath==''):
         Errormessage('Enter File Path')   
    elif(filename==''):
         Errormessage('File name cannot be blank')
    elif(is_valid_filename(filename)!=True):
         Errormessage('Invalid file name')
    elif(combo.get()=='Excel'):
         createExcel()
    elif(combo.get()=='Word'):
         createWord()
    else:
         Errormessage('Select File Type')

 
def saveToexcel(img,imgname):
    filename=e.get()
    filepath=str(z.get()).replace("/","\\\\")
    if(path.exists(filepath+'\\'+filename+'.xlsx')):
         try:
               wb=load_workbook(filepath+'\\'+filename+'.xlsx')
               ws = wb.worksheets[0]
               sheet=wb.active
               maxrow=sheet.max_row
               offset=maxrow+1
               commentPtr = str('A'+str(int(offset-1)))
               if(Comment.get()=='Screenshot Description'):
                    sheet[commentPtr] = ''
               else:
                    sheet[commentPtr] = Comment.get()
               eof = str('A'+str(int(offset+40))) 
               sheet[eof] = "..."
               img.save(imgname,'PNG') 
               img = openpyxl.drawing.image.Image(imgname)
               img.anchor = 'A'+str(offset)
               ws.add_image(img)
               wb.save(filepath+'\\'+filename+'.xlsx')
               entry_text1.set('Screenshot Description')
               
         except PermissionError:
              Errormessage('File currently open.Close to Continue')
    else:
         Errormessage('No such file in path')

def saveToWord(img,imgname):
    filename=e.get()
    filepath=str(z.get()).replace("/","\\\\")
    if(path.exists(filepath+'\\'+filename+'.docx')):
         try:
              document = Document(filepath+'\\'+filename+'.docx')
              if(Comment.get()=='Screenshot Description'):
                   document.add_paragraph('')
              else:
                   document.add_paragraph(Comment.get())
              img.save(imgname,'PNG') 
              img = document.add_picture(imgname, width=Inches(7))
              document.save(filepath+'\\'+filename+'.docx')
              entry_text1.set('Screenshot Description')
         except PermissionError:
              Errormessage('File currently open.Close to Continue')
    else:
         Errormessage('No such file name in path')


def grabScreenShot():
    filename=e.get()
    filepath=str(z.get()).replace("/","\\\\")
    if(filepath==''):
        Errormessage('Enter File Path')        
    elif(filename==''):
         Errormessage('File name cannot be blank')
    elif(is_valid_filename(filename)!=True):
         Errormessage('Invalid file name')
    elif(combo.get()=='Excel'):
          root.wm_state('iconic')
          time.sleep(0.5)
          im = ImageGrab.grab() 
          root.wm_state('normal')
          saveToexcel(im,filepath+'\\lastcopiedimg.png')
    elif(combo.get()=='Word'):
          root.wm_state('iconic')
          time.sleep(0.5)
          im = ImageGrab.grab() 
          root.wm_state('normal')
          saveToWord(im,filepath+'\\lastcopiedimg.png')
    else:
          Errormessage('Select File Type')
   
def grabClipboard():
    filename=e.get()
    filepath=str(z.get()).replace("/","\\\\")
    if(filepath==''):
        Errormessage('Enter File Path')        
    elif(filename==''):
         Errormessage('File name cannot be blank')
    elif(is_valid_filename(filename)!=True):
         Errormessage('Invalid file name')
    elif(combo.get()=='Excel'):
          root.wm_state('iconic')
          time.sleep(0.5)
          im = ImageGrab.grabclipboard()
          root.wm_state('normal')
          saveToexcel(im,filepath+'\\lastcopiedimg.png')
    elif(combo.get()=='Word'):
          root.wm_state('iconic')
          time.sleep(0.5)
          im = ImageGrab.grabclipboard()
          root.wm_state('normal')
          saveToWord(im,filepath+'\\lastcopiedimg.png')
    else:
          Errormessage('Select File Type')
    
def search_for_file_path ():
    currdir = os.getcwd()
    tempdir = filedialog.askdirectory( parent=root,initialdir=currdir, title='Please select a directory')
    entry_text.set(tempdir)
    e.delete(0, END)

frame = ttk.Frame(root)
frame.grid(column=1, row=1)
root.title("SnapIT")
# root.iconbitmap(r"img.ico") 
cpane = CollapsiblePane(root,'\t\t File Settings \t\t'+u"\u23EB", '\t\t File Settings \t\t'+ u"\u23EC") 
cpane.grid(row = 0, columnspan=2,sticky='nwes')

root.resizable(False,False)
selectPath = ttk.Button(cpane.frame, text='Set File Location', width = 40,command = search_for_file_path)
selectPath.grid(row=0,columnspan=2,sticky='nwes')
entry_text = StringVar()
z = ttk.Entry(cpane.frame,width=40,textvariable=entry_text)
z.grid(row=1,columnspan=2,sticky='nwes')
z.config(state='readonly')

ttk.Label(cpane.frame,text = 'Enter File Name : ',anchor='center',width=20).grid(row=2,sticky='nwes')
def clear_search(event):
#     e.delete(0, END)
    e.focus_force()
e = ttk.Entry(cpane.frame,width=20)
e.grid(row=2,column=1,sticky='nwes')
e.bind("<Button-1>", clear_search)

ttk.Label(cpane.frame,text = 'File Type : ',anchor='center',width=20).grid(row=3,sticky='nwes')
variable = StringVar()
variable.set("Select")
combo = ttk.Combobox(cpane.frame ,width=20,textvariable=variable,values=["Excel","Word"])
combo.grid(row=3,column=1,sticky='nwes')
combo.config(state='readonly')

Create = ttk.Button(cpane.frame, text='Create File', width = 40,command=CreateFile)
Create.grid(row=4,columnspan =2 ,sticky='nwes')

ScreenShot = ttk.Button(root,text='Screenshot',width=20,command=grabScreenShot)
ScreenShot.grid(row=5,sticky='nwes')

ClipBoard = ttk.Button(root,text='Clipboard',width=20,command=grabClipboard)
ClipBoard.grid(row=5,column=1,sticky='nwes')

entry_text1 = StringVar()
entry_text1.set('Screenshot Description')
def clear_search_Comment(event):
    Comment.delete(0, END)
    Comment.focus_force()
Comment = ttk.Entry(root,width=40,textvariable=entry_text1)
Comment.grid(row=6,columnspan=2,sticky='nwes')
Comment.bind("<Button-1>", clear_search_Comment)
root.mainloop()
